#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成嵌入式竞赛报告初稿 (基于模版结构)"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

d = Document()

# 中文默认字体
d.styles['Normal'].font.name = '宋体'
d.styles['Normal'].font.size = Pt(10.5)

def h1(t): d.add_heading(t, level=1)
def h2(t): d.add_heading(t, level=2)
def h3(t): d.add_heading(t, level=3)
def p(t):
    para = d.add_paragraph(t)
    return para
def img(t):
    para = d.add_paragraph()
    run = para.add_run("【此处插入图片：" + t + "】")
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

# ── 标题 ──
title = d.add_paragraph()
tr = title.add_run("基于 RK3588 NPU 的暗光图像实时增强与车牌识别系统")
tr.bold = True; tr.font.size = Pt(18)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
d.add_paragraph()

# ── 摘要 ──
h2("摘要")
p("夜间及低照度环境下，普通摄像头成像噪声大、细节丢失，车牌等关键信息难以辨认，"
  "给智能交通与安防监控带来困难。本作品基于飞凌 ELF2（瑞芯微 RK3588）嵌入式开发板，"
  "设计并实现了一套完全运行在边缘端的“暗光图像实时增强 + 车牌识别”系统。"
  "系统以 ZWO ASI585MC（Sony IMX585）RAW 相机为输入，采用长短交替曝光采集三帧 RAW 图像，"
  "送入自研的 NBINet 多帧融合去噪网络，在 RK3588 内置 NPU 上以 FP16 精度实时推理，"
  "输出清晰的 RGB 图像；去噪结果一路经 DRM/KMS 直接显示到 MIPI 屏，"
  "一路送入 HyperLPR3 两阶段车牌识别（YOLO 检测 + CTC 序列识别），"
  "识别出的车牌号、置信度、抓拍图写入本地，并通过板载 WiFi 提供的轻量 HTTP 服务"
  "供手机浏览器实时查询。整套推理流水线以多线程 C 语言实现（采集、去噪、显示、识别"
  "四线程 + 环形缓冲），端到端约 8fps，无需联网、无需 PC，插电即用。"
  "实测中系统可将暗光噪声图像增强为清晰彩色图，并正确识别车牌（如“苏A6D8F8”），"
  "手机端可按车牌归组查询抓拍照片与命中帧数。本作品完整覆盖"
  "“RAW 采集 → NPU 去噪 → 车牌识别 → 显示与查询”的端侧闭环，"
  "具有实时、低功耗、可离线部署的特点。")

# ── 第一部分 ──
h1("第一部分  作品概述")

h2("1.1  功能与特性")
p("① 暗光实时增强：采用 RAW 域长短三帧融合去噪，将低照度噪声图像增强为清晰彩色图像，"
  "保留暗部细节与真实色彩。\n"
  "② 车牌检测与识别：对增强后图像进行车牌定位（YOLO）与字符识别（CTC），"
  "支持中国蓝牌/黄牌等，输出车牌号与置信度。\n"
  "③ 实时显示：去噪结果经 DRM/KMS 直接输出到 MIPI DSI 屏（1024×600），画面无撕裂。\n"
  "④ 手机查询：板载 WiFi 提供 HTTP 服务，手机浏览器即可查询识别记录，"
  "按车牌自动归组、显示抓拍图与命中帧数，支持车牌号搜索。\n"
  "⑤ 离线自治：全流程运行于 RK3588 边缘端，无需云端与 PC，上电自动运行。")
img("系统实时增强前后对比 / 手机查询界面")

h2("1.2  应用领域")
p("本作品属于赛题“选题方向二：端侧 AI 视觉应用”（智能安防方向），"
  "面向夜间及低照度场景的车辆管理与安防需求。典型应用包括：智慧停车场出入口"
  "车牌识别（以算法增强替代物理补光）；小区、园区门禁车辆放行；夜间道路卡口、违停抓拍；"
  "执法记录仪与车载巡逻终端；无人值守站点的车辆通行记录。相较于传统“高功率补光灯 + "
  "普通摄像头”方案，本系统以算法增强替代物理补光，降低能耗与光污染，且边缘端离线运行，"
  "适合供电与网络受限的野外或移动场景。作品紧扣该方向的核心考核点——"
  "特定光照条件下的图像清晰度与色彩还原、目标检测帧率与识别准确率、"
  "以及复杂光照下的系统鲁棒性与长时间稳定性。")
img("应用场景示意")

h2("1.3  主要技术特点")
p("① RAW 域“短-长-短”三帧融合去噪：去噪网络 NBINet 借鉴 D2HNet 的三帧融合思路，"
  "直接在传感器 12-bit RAW Bayer 域端到端处理——以短曝光帧保留清晰边缘（运动锚点）、"
  "长曝光帧补足暗部，融合出高信噪比图像，绕过传统 ISP，效果优于普通 RGB 视频后处理。\n"
  "② NPU FP16 端侧加速：NBINet 经 PyTorch 训练 → ONNX → RKNN 转换后部署于 RK3588 NPU。"
  "实测发现 INT8 量化会破坏去噪模型的颜色重建，最终采用 FP16，在精度与速度间取得平衡。\n"
  "③ NPU+CPU 异构并行：去噪模型运行于 NPU，车牌识别（HyperLPR3：YOLO 检测 + CTC 序列识别 "
  "+ 单/双层分类）经 ncnn 在 CPU（ARM Cortex-A76/A55，NEON 加速）上推理；两者分居 NPU 与 CPU、"
  "在不同线程并行执行，NPU 去噪与 CPU 识别同时进行、互不争抢，充分利用 RK3588 的异构算力。\n"
  "④ 全 C 多线程实时流水线：采集、去噪、显示、识别四线程，以环形缓冲解耦，"
  "去噪与识别各自消费全部帧，互不阻塞。\n"
  "⑤ 显示与查询：DRM/KMS 直控 MIPI 屏；自研无第三方依赖的 C HTTP 服务，手机免装 APP 即可查询。")

h2("1.4  主要性能指标")
tbl = d.add_table(rows=1, cols=2)
tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = "指标"; tbl.rows[0].cells[1].text = "参数"
metrics = [
    ("主控平台", "瑞芯微 RK3588（4×A76 + 4×A55，6 TOPS NPU）"),
    ("图像传感器", "ZWO ASI585MC（Sony IMX585，12-bit RAW Bayer）"),
    ("去噪模型", "NBINet（基于 D2HNet 改造），三帧 RAW 融合，FP16"),
    ("去噪精度", "验证集 PSNR 35.0 dB，SSIM 0.924"),
    ("去噪分辨率", "输入 272×480 → 输出 544×960 RGB"),
    ("去噪速度", "约 120 ms/帧（约 8 fps，NPU）"),
    ("车牌识别耗时", "约 40–60 ms/帧（ncnn，CPU ARM NEON，与 NPU 去噪并行）"),
    ("显示", "MIPI DSI 1024×600，DRM/KMS 直显"),
    ("端到端帧率", "约 8 fps，实测无丢帧"),
    ("联网查询", "板载 WiFi + HTTP 服务，手机浏览器"),
]
for k, v in metrics:
    row = tbl.add_row().cells
    row[0].text = k; row[1].text = v

h2("1.5  主要创新点")
p("【算法层】\n"
  "① RAW 域“短-长-短”三帧融合去噪：去噪网络 NBINet 借鉴 D2HNet 的三帧融合思路，"
  "直接在 12-bit RAW 域端到端完成多帧对齐、融合与增强——以两帧短曝光作为运动“锚点”"
  "保留清晰边缘、长曝光帧补足暗部亮度，绕过传统 ISP，从底层同时抑制噪声与运动拖影；\n"
  "② 面向 NPU 的算子级改造（核心创新）：将原网络中 NPU 不支持或低效的算子替换为 NPU 友好算子——"
  "下采样由离散小波变换（DWT）改为二维卷积（Conv2d）、上采样由小波反变换（IWT）改为像素重组"
  "（PixelShuffle）、并对对齐模块（DCN）做 NPU 友好化改造，在几乎不损失效果的前提下，"
  "使一个面向 GPU 的研究网络得以在 RK3588 NPU 上原生高效加速；\n"
  "【工程/部署层】\n"
  "③ 端侧 FP16 部署，规避 INT8 量化导致的去噪模型色彩失真与分辨率退化；\n"
  "④ NPU+CPU 异构并行流水线：去噪运行于 NPU、车牌识别经 ncnn 运行于 CPU，双环形缓冲、并行执行，"
  "每帧兼顾显示与识别；\n"
  "⑤ 边缘自治 + 手机免装 APP 查询：板端离线自治，板载 WiFi + 轻量 HTTP 服务支持手机按车牌归组查询。")

h2("1.6  设计流程")
p("夜间 RAW 数据采集 → PyTorch 训练 NBINet 去噪网络 → 导出 ONNX → RKNN 工具链转换为 "
  "NPU 可执行模型（FP16）→ C 语言集成相机采集、NPU 推理、DRM 显示、ncnn 车牌识别 → "
  "ELF2 板端部署，联调解决 USB 带宽、模型量化、图像朝向等问题 → 增加 WiFi 与手机查询 → "
  "系统联测与优化。")
img("设计流程图")

# ── 第二部分 ──
h1("第二部分  系统组成及功能说明")

h2("2.1  整体介绍")
p("系统整体采用“生产者—消费者”多线程流水线，各模块以环形缓冲解耦。数据流为："
  "相机输出 RAW → 采集线程（Bayer 提取为 4 通道 RGGB 浮点）→ 原始环形缓冲 → "
  "去噪线程（取短-长-短三帧融合，RKNN NPU 推理）→ 去噪 RGB 分发到两路："
  "① 显示环形缓冲 → 显示线程（DRM/KMS 直显 MIPI 屏）；"
  "② 识别环形缓冲 → 识别线程（YOLO 检测 + CTC 识别 → 写入记录与抓拍图）→ "
  "HTTP 服务 → 手机查询。四线程并行、互不阻塞，保证每一帧既显示又识别。"
  "其中去噪运行于 NPU、车牌识别经 ncnn 运行于 CPU，二者并行，形成 NPU+CPU 异构协同。")
img("系统整体框图（各子模块与数据流）")

h2("2.2  硬件系统介绍")
h3("2.2.1  硬件整体介绍")
p("硬件由 ELF2（RK3588）核心板、ZWO ASI585MC USB 相机、MIPI DSI 显示屏（1024×600）、"
  "板载 CF-AX200 WiFi/蓝牙模块构成。相机经 USB 接入，MIPI DSI 输出至显示屏，"
  "WiFi 提供手机查询网络。")
img("硬件连接总图")
h3("2.2.2  机械设计介绍")
p("本作品以开发板与工业相机为核心，未做专门机械结构设计；实际部署时可加装外壳，"
  "固定相机与屏幕并做防护。")
img("整机布置照片 / 结构示意（如有）")
h3("2.2.3  电路各模块介绍")
p("主要接口：5V 供电、USB 相机接口、MIPI DSI 屏接口。需要说明的是，"
  "RK3588 的 USB3（xHCI/DWC3）在持续高带宽流传输时存在稳定性问题，"
  "本作品经对照测试后选用 USB2 通道并配合相机硬件 ROI 裁剪降低带宽，保证采集稳定。")
img("接口/供电示意图")

h2("2.3  软件系统介绍")
h3("2.3.1  软件整体介绍")
p("软件分三部分：① PC/服务器端——模型训练（PyTorch）与 RKNN 模型转换；"
  "② 板端——C 语言实时推理流水线（四线程）；③ 手机端——浏览器访问板端 HTTP 服务查询。"
  "三者中，板端为核心，独立离线运行。")
img("软件层次结构图")
h3("2.3.2  软件各模块介绍")
p("• 采集模块（asi_capture）：相机 2×2 binning 并设硬件 ROI 直接输出中心 960×544 区域，"
  "长短曝光交替；将 RAW Bayer 按 RGGB 提取为 4 通道浮点，写入原始环形缓冲。\n"
  "• 去噪模块（npu_worker）：从缓冲取“短-长-短”三帧，打包为 NPU 双输入"
  "（short_cat 8 通道 + long_raw 4 通道，FP16），RKNN 推理输出 RGB，"
  "并克隆分发到显示与识别两个环形缓冲。\n"
  "• 显示模块（fb_display）：以 DRM/KMS 建立 dumb buffer 并 modeset，"
  "将去噪图双线性缩放至 1024×600，经对比度/饱和度增强后直写扫描缓冲显示。\n"
  "• 识别模块（lpr_worker）：去噪图缩放至 320×320 送 YOLO 检测 → NMS → 取最高分框裁切 → "
  "CTC 识别车牌字符 → 单/双层分类 → 写入 plates.txt 与抓拍图。该模块经 ncnn 运行于 CPU"
  "（ARM Cortex-A76/A55，NEON 加速），与 NPU 上的去噪线程并行，实现 NPU+CPU 异构协同。\n"
  "• Web 服务（lpr_server）：基于 socket 的无依赖 C HTTP 服务，读取记录按车牌归组，"
  "返回 JSON 与网页，并提供抓拍图访问。")
img("各模块流程图（采集/去噪/识别）")

h3("2.3.3  数据集采集与模型训练")
p("【本小节数据采集部分由数据采集同学补充，以下为框架建议，请据实填写具体内容】")
p("• 采集设备与场景：使用 ASI585MC（IMX585）RAW 相机，在夜间及低照度真实场景采集；\n"
  "• 数据规模：共采集约 3000–4000 组样本；\n"
  "• 单组内容：长/短交替曝光的多帧 RAW（短1 / 长 / 短2）及多帧平均得到的低噪声真值（GT）；\n"
  "• 采集条件（待补充）：具体地点、时间、光照强度、天气、车辆与车牌分布等；\n"
  "• 数据处理：RAW Bayer 按 RGGB 提取为 4 通道、黑电平校正与归一化，GT 经去马赛克转 RGB；\n"
  "• 数据划分（待补充）：训练/验证/测试集比例及划分方式；\n"
  "• 数据清洗与标注（待补充）：无效帧剔除、对齐、（如有）车牌区域标注等流程。")
p("模型与训练：去噪网络 NBINet 借鉴 D2HNet 的“短-长-短”三帧 RAW 融合框架，"
  "输入为短曝光拼接（8 通道）+ 长曝光（4 通道），端到端输出去噪 RGB。"
  "为适配 RK3588 NPU，对原网络中 NPU 不支持或低效的算子进行了替换：下采样由离散小波变换（DWT）"
  "改为二维卷积（Conv2d）、上采样由小波反变换（IWT）改为像素重组（PixelShuffle），"
  "并对对齐模块（DCN）做 NPU 友好化改造，使模型可被 NPU 原生高效加速。"
  "以低噪声 GT 为监督用 PyTorch 训练（损失函数、训练轮数等待补充），"
  "训练完成后导出 ONNX，再经 RKNN 工具链转换为 RK3588 NPU 可执行的 FP16 模型完成板端部署。")
img("数据采集场景 / 数据集样例（长短曝光 RAW 与 GT）")
img("模型训练收敛曲线 / 去噪效果验证对比")

# ── 第三部分 ──
h1("第三部分  完成情况及性能参数")

h2("3.1  整体介绍")
p("系统已完成整体联调，实现从相机采集到手机查询的端侧闭环，上电自动运行。")
img("整机正面 / 斜 45° 全局照片")

h2("3.2  工程成果")
h3("3.2.1  机械成果")
img("整机布置实物照片")
h3("3.2.2  电路成果")
img("板卡与接口连接实物照片")
h3("3.2.3  软件成果")
p("屏幕实时显示去噪后清晰彩色画面；手机浏览器访问板端服务，可见按车牌归组的识别记录、"
  "抓拍图廊与搜索框。")
img("屏幕去噪画面照片")
img("手机端车牌查询界面截图")

h2("3.3  特性成果")
p("对照赛题选题方向二的考核指标，本作品实现结果如下：\n"
  "① 图像质量（图像预处理）：在低照度条件下，将高噪声暗光图像增强为清晰彩色图像，"
  "暗部细节与真实色彩得到恢复，色彩还原准确（去噪输出通道顺序与训练一致）；"
  "在验证集上达到 PSNR 35.0 dB、SSIM 0.924，重建质量与结构相似度均较高。\n"
  "② 视觉算法性能：车牌检测 + 识别端到端约 8 fps；实测正确识别车牌“苏A6D8F8”，"
  "检测置信度约 0.8，去噪约 120 ms/帧、识别约 40–60 ms/帧。\n"
  "③ 系统鲁棒性：识别在车牌远近、角度变化下自适应，"
  "对倒装/镜像等成像方向问题已在流水线中校正；识别置信度低于阈值时不误报。\n"
  "④ 系统稳定性：四线程流水线长时间运行无丢帧（drops=0），"
  "采集通道稳定（规避 USB3 主控稳定性问题），数据写入与手机查询稳定。\n"
  "⑤ 手机查询：识别记录按车牌自动归组，显示全部抓拍图、命中帧数与置信度，支持车牌号搜索。")
img("去噪前后对比图")
img("车牌识别结果截图（苏A6D8F8）")

# ── 第四部分 ──
h1("第四部分  总结")

h2("4.1  可扩展之处")
p("① 修复 RK3588 USB3（xHCI/DWC3）驱动后，可提升相机带宽与帧率，支持更高分辨率；"
  "② 采用滑窗帧复用（将上一组的后短曝光帧复用为下一组的前短曝光帧），"
  "由“3 帧新帧出 1 帧”变为“2 帧新帧出 1 帧”，在不增加算力的前提下提升有效帧率；"
  "③ 支持多路相机与多车牌并行识别；④ 识别结果可上云或推送，形成车辆通行数据库；"
  "⑤ 采用 Realtek 芯片 USB 网卡或独立路由，提升手机接入的稳定性；"
  "⑥ 采用混合量化（敏感层 FP16 + 其余 INT8）进一步提速；"
  "⑦ 增加车型/车色识别、真实时间戳（需 RTC 备用电池）等功能。")

h2("4.2  心得体会")
p("本作品的最大收获，是完整走通了一条“训练侧模型 → 边缘端实时系统”的工程链路，"
  "并在过程中解决了大量真实的底层问题。\n\n"
  "第一，硬件与驱动层面的坑往往最隐蔽。相机在 USB3 口持续采集时频繁复位、无法出图，"
  "我们通过“最简单帧采集 vs 持续流采集”“USB3 vs USB2”的对照实验，"
  "定位到是 RK3588 的 xHCI 控制器在 SuperSpeed 持续流下的稳定性问题，"
  "最终以 USB2 通道 + 相机硬件 ROI 裁剪（降低带宽的同时正好取到模型所需区域）绕开，"
  "把帧率从 2fps 提升到 8fps。这让我们体会到：边缘部署中，先用最小实验隔离问题、"
  "再对症下药，比盲目改代码高效得多。\n\n"
  "第二，模型量化并非“无损搬运”。最初 INT8 量化后的去噪结果出现严重偏色（品红、绿弱），"
  "我们通过逐环节 dump 中间数据、在 PC 上用 ONNX 运行时逐层比对，"
  "确认是 INT8 量化破坏了去噪网络对通道细微差异的重建（并使上采样层失效、分辨率减半），"
  "改用 FP16 后颜色与分辨率均恢复正常。对“输出即图像”的重建类模型，精度优先于速度。\n\n"
  "第三，端到端系统里每一层的约定都要对齐。车牌识别一度输出乱码，"
  "排查后发现是多个问题叠加：相机成像 180° 倒装导致字符方向错误、"
  "YOLO 输入归一化误除 255 使图像接近全黑、CTC 时间步数写错导致读越界产生固定乱码尾巴。"
  "逐一修正后才正确识别出车牌。这让我们深刻理解到：预处理、通道顺序、张量形状等“接口约定”，"
  "任何一处不一致都会让整条链路失效，必须用数据验证而非想当然。\n\n"
  "此外，我们还实践了 DRM/KMS 直接显示、无第三方依赖的 C HTTP 服务、"
  "多线程环形缓冲流水线等工程技术。整个过程锻炼了从算法、系统到硬件的全栈调试能力，"
  "也让我们认识到，一个能实际运行的边缘 AI 系统，其价值不仅在于模型本身，"
  "更在于把每一层可靠地拼接起来。")

# ── 第五部分 ──
h1("第五部分  参考文献")
refs = [
    "ZHAO Y, XU Y, YAN Q, et al. D2HNet: joint denoising and deblurring with "
        "hierarchical network for robust night image restoration[C]//European "
        "Conference on Computer Vision (ECCV). Cham: Springer, 2022.",
    "GRAVES A, FERNÁNDEZ S, GOMEZ F, et al. Connectionist temporal classification: "
        "labelling unsegmented sequence data with recurrent neural networks[C]//"
        "Proceedings of the 23rd International Conference on Machine Learning (ICML). "
        "New York: ACM, 2006: 369-376.",
    "MILDENHALL B, BARRON J T, CHEN J, et al. Burst denoising with kernel prediction "
        "networks[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern "
        "Recognition (CVPR). Piscataway: IEEE, 2018: 2502-2510.",
    "JOCHER G. YOLOv5[EB/OL]. (2020)[2026-07-07]. https://github.com/ultralytics/yolov5.",
    "HyperLPR: 高性能中文车牌识别框架[EB/OL]. [2026-07-07]. "
        "https://github.com/szad670401/HyperLPR.",
    "腾讯. ncnn: 高性能神经网络前向计算框架[EB/OL]. [2026-07-07]. "
        "https://github.com/Tencent/ncnn.",
    "Rockchip. RKNN-Toolkit2[EB/OL]. [2026-07-07]. "
        "https://github.com/airockchip/rknn-toolkit2.",
    "The Linux Kernel Community. Kernel mode setting (KMS)[EB/OL]. [2026-07-07]. "
        "https://www.kernel.org/doc/html/latest/gpu/drm-kms.html.",
    "瑞芯微电子股份有限公司. RK3588 技术参考手册[Z]. 福州: 瑞芯微电子, 2023.",
    "Sony Semiconductor Solutions. IMX585 CMOS image sensor datasheet[Z]. Tokyo: "
        "Sony, 2022.",
    "保定飞凌嵌入式技术有限公司. ELF2 (RK3588) 开发板快速使用手册[Z]. 保定: 飞凌嵌入式, 2024.",
]
for i, r in enumerate(refs, 1):
    d.add_paragraph("[%d] %s" % (i, r))

out = "/mnt/c/Users/HP/Desktop/嵌赛报告_初稿_v7.docx"
d.save(out)
print("已生成:", out)
print("段落数:", len(d.paragraphs), " 表格数:", len(d.tables))
